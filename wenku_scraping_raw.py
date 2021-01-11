from selenium import webdriver
from bs4 import BeautifulSoup as BS
from selenium import webdriver
import time
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
import os
from PIL import Image

get_url = input('Please input the url: ')

file_size = input('please input the amount of pages: ')

file_name = input('Please input the file name : ')

module_name = input("File's type: ")

url = get_url
driver = webdriver.Chrome()
driver.get(url)
time.sleep(int(file_size) * 10) #experiment with timer to fetch all the data
page = driver.page_source
driver.quit()
soup = BS(page, 'html.parser')


def pdf_img():
    files = soup.find_all('div', attrs={'class': 'ppt-image-wrap'})
    raw_links = [file.find('img') for file in files]
    links = []
    for link in raw_links:
        try:
            links.append(link['src'])
        except:
            links.append(link['data-src'])
    print(links)
    i = 1
    os.mkdir(file_name)
    for link in links:
        img_data = requests.get(link).content
        with open(f'{os.getcwd()}/{file_name}/{i}.jpg', 'wb') as handler:
            handler.write(img_data)
        i += 1
    path = f'{os.getcwd()}/{file_name}/'
    demo = Image.open(fr'{path}/1.jpg').convert('RGB')
    list1 = []
    for a in range(i - 1):
        a = Image.open(fr"{path}{a + 1}.jpg").convert('RGB')
        list1.append(a)
        demo.save(fr'{path}images.pdf', save_all=True, append_images=list1)


def docx():
    files = soup.find_all('div', attrs={'class': 'reader-txt-layer'})
    texts = [file.text.strip() for file in files]
    print(texts)
    return texts


def txt():
    files = soup.find_all('p', attrs={'class': 'p-txt'})
    texts = [file.text.strip() for file in files]
    print(texts)
    return texts


def pdf_txt():
    files = soup.find_all('p', attrs={'class': 'reader-word-layer'})
    texts = [file.text.strip() for file in files]
    print(texts)
    return texts


def file_packing():
    doc = Document()
    run = doc.add_paragraph().add_run()
    font = run.font
    font.name = 'SimSun-ExtB 常规'  #字体种类
    paragraph = doc.add_paragraph()
    for text in eval(module_name)():
        paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ph_format = paragraph.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(12)  # 设置段后间距
    ph_format.first_line_indent = Inches(0.5)
    ph_format.line_spacing = Pt(19)
    doc.save(f'{file_name}.docx')


if module_name == 'pdf_img':
    pdf_img()
else:
    file_packing()


